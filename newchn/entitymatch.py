import re
import docx
from glob import glob
import os
from sim_simhash import match
#s = '甲方：中国移通信集团江苏有限公司无锡分公司 \n 地址：无锡吴518号   无锡吴都路578号\n 负责人： 谢比伯  射生新,   中国移通信集团江苏有限'

def write_result(doc, all_error):

    doc = os.path.splitext(os.path.split(doc)[1])[0]

    with open('./test_result/entity_cpy.txt', 'a') as f:

        #dict = zip(para_no, doc_err_comp, doc_err_add, doc_err_sign)
        '''
        for fn, ft  in dict:
            ft_str = ''.join(str(ft).replace('[','',1).rsplit(']',1))
            f.write(doc + ' | ' + fn + ' | ' + ft_str + ' | ' + rw + ' | ' + str(m) + '\n')
        '''
        allerr_str = ''.join(str(all_error).replace('[','',1).rsplit(']',1))
        if all_error[0] ==[] and all_error[1] == [] and all_error[2] == []:
            f.write(doc + '|' + allerr_str + '|' + '一致' + '\n')
        else:
            f.write(doc + '|' + allerr_str + '|' + '不一致' + '\n')


company_name = '中国移动通信集团江苏有限公司'
address = '无锡市吴都路518号'
add_old = '无锡市滨湖区蠡园经济开发区溢红路139号'
sign_name = '谢生勃'

#comp_dict = dict(zip(company_name))
#add_dict = dict(zip(address))
#sign_dict = dict(zip(sign_name))

#doc_path = glob('./sim/*.*docx')
#doc_path = glob('/home/eric/aiprogram/newchn/contract_zxj/weituojiamenglei/wtjm2/*.*docx')

def entity_match(doc_path):
    for doc in sorted(doc_path):
        try:
            file = docx.Document(doc)
            print('The Doc name is {}'.format(doc))
        except :
            print('*******The File {} Is No Support To Read!*********'.format(doc))
            continue
        doc_err_comp = []
        doc_err_add = []
        doc_err_sign = []
        all_error =[]
        para_index = 0
        p_index = 0
        for para in file.paragraphs:
            if p_index <=2:
                #print('title is :{}'.format(para.text))
                p_index += 1
                continue

            name_list =[]
            name_index =[]

            comp_str = re.findall(r'中国.*?司', para.text)
            add_str = re.findall(r'无.*?号', para.text)
            para_index += 1
            #print(comp_str)
            #print(add_str)

            for item in comp_str:
                sim_comp = 0
                redundant = 0
                for word in item:
                    if word in company_name:
                        sim_comp += 1
                if len(item) > len(company_name):
                    redundant = len(item)-len(company_name)
                sim  = (sim_comp - redundant) / len(company_name)
                print('company_name item {} sim: {}'.format(item,sim))
                #sim = match(item, company_name)
                #print('company name item {} sim: {}'.format(item,sim))
                #if item != company_name and len(item) <17 and sim > 0.6 :
                if item != company_name and len(item) <17 and sim > 0.88:
                    doc_err_comp.append(item)


            for item in add_str:
                sim_add = 0
                for word in item:
                    if word in address:
                        sim_add += 1
                sim  = sim_add / len(address)
                print('address item {} sim: {}'.format(item,sim))
                #sim = match(item, address)
                #print('address item {} sim: {}'.format(item,sim))
                #if item != address and len(item) < 12 and sim > 0.6:
                if item != address and len(item) < 12 and sim > 0.7 or item == add_old:
                    doc_err_add.append(item)


            name = re.finditer(r'[谢射]', para.text)

            for i in name:
                #print(i)
                name = i.span()
                name_index.append(name[0])

            for n in name_index:
                name_list.append(para.text[n:n+3])

            #print(name_list)

            for item in name_list:
                #sim = match(item, sign_name)
                sim_sign = 0
                for word in item:
                    if word in sign_name:
                        sim_sign += 1
                sim  = sim_sign / len(sign_name)
                print('sign_name item {} sim: {}'.format(item,sim))
                if item != sign_name and sim > 0.5:
                    doc_err_sign.append(item)


        all_error.append(doc_err_comp)
        all_error.append(doc_err_add)
        all_error.append(doc_err_sign)

        print(all_error)
        write_result(doc,  all_error)

        #print('error company title: {} \nerror address:{} \nerror sign name:{} \n'.format(error_comp_str, error_add_str,doc_err_sign))

file_list = []
def readoc(doc_path):
    doc_list = os.listdir(doc_path)
    #print(doc_list)
    for doc in doc_list:
        doc = os.path.join(doc_path, doc)
        if os.path.isdir(doc):
            readoc(doc)
        else:
            file_list.append(doc)

    return file_list

if __name__=='__main__':
    #file_list = readoc('/home/eric/aiprogram/newchn/contract_zxj/weituojiamenglei/wtjm2')
    #file_list = readoc('/home/eric/aiprogram/newchn/contract_zxj/wuxihetong/wxht1
    file_list = readoc('/home/eric/aiprogram/newchn/sim')
    entity_match(file_list)

'''
doc_path = glob('/home/eric/aiprogram/test_doc/*.*')

for doc in sorted(doc_path):
    file = docx.Document(doc)
    print('The Doc name is {}'.format(doc))

    for para in file.paragraphs:

        comp_str = re.findall(company_name, para)
'''
